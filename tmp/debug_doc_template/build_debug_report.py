from __future__ import annotations

import csv
import hashlib
import shutil
import zipfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path("/Users/riku/Project/program/1week_project/AI_judge")
REFERENCE = Path(
    "/Users/riku/.codex/plugins/cache/openai-curated-remote/openai-templates/"
    "0.1.0/skills/artifact-template-experiment-analysis/assets/reference.docx"
)
OUT = ROOT / "AI_judge_debug_progress_report.docx"


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def set_run_font_for_japanese(run) -> None:
    font_name = "Arial Unicode MS"
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = rpr._new_rFonts()
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def set_paragraph_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        run = paragraph.add_run()
    else:
        run = paragraph.runs[0]
    run.text = text
    set_run_font_for_japanese(run)
    for extra in list(paragraph.runs[1:]):
        extra._element.getparent().remove(extra._element)


def set_cell_text(cell, text: str) -> None:
    if cell.paragraphs:
        set_paragraph_text(cell.paragraphs[0], text)
        for p in cell.paragraphs[1:]:
            p._element.getparent().remove(p._element)
    else:
        set_paragraph_text(cell.add_paragraph(), text)


def fill_table(table, rows: list[list[str]]) -> None:
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            if r < len(table.rows) and c < len(table.columns):
                set_cell_text(table.cell(r, c), text)


def read_loss_summary() -> dict[str, str]:
    loss_path = ROOT / "loss.csv"
    if not loss_path.exists():
        return {
            "best_epoch": "要確認",
            "best_valid_accuracy": "要確認",
            "final_valid_accuracy": "要確認",
            "final_train_accuracy": "要確認",
            "final_valid_loss": "要確認",
        }

    rows = []
    with loss_path.open(newline="", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            rows.append(row)
    if not rows:
        return {
            "best_epoch": "要確認",
            "best_valid_accuracy": "要確認",
            "final_valid_accuracy": "要確認",
            "final_train_accuracy": "要確認",
            "final_valid_loss": "要確認",
        }

    best = max(rows, key=lambda r: float(r["valid_accuracy"]))
    final = rows[-1]
    return {
        "best_epoch": best["epoch"],
        "best_valid_accuracy": f'{float(best["valid_accuracy"]) * 100:.2f}%',
        "final_valid_accuracy": f'{float(final["valid_accuracy"]) * 100:.2f}%',
        "final_train_accuracy": f'{float(final["train_accuracy"]) * 100:.2f}%',
        "final_valid_loss": f'{float(final["valid_loss"]):.4f}',
    }


def replace_package_text(docx_path: Path, old: str, new: str) -> None:
    tmp_path = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.endswith(".xml"):
                    data = data.replace(old.encode("utf-8"), new.encode("utf-8"))
                zout.writestr(item, data)
    tmp_path.replace(docx_path)


def main() -> None:
    reference_hash = sha256(REFERENCE)
    shutil.copyfile(REFERENCE, OUT)

    doc = Document(OUT)
    today = date.today().isoformat()
    loss = read_loss_summary()

    replacements = {
        0: "Debug Progress Report",
        6: "AI_judge",
        7: "Debug Progress Report",
        23: "Riku / Codex draft",
        24: today,
        26: "Document Control",
        29: "Purpose",
        30: "This document is a first-pass debug progress report for AI_judge. It summarizes what appears to be complete, what is unstable, and which fixes should be considered first.",
        31: "Detailed logs and screenshots can be added later. This draft uses only the README, visible source files, training log, and current working-tree state.",
        33: "Business Context",
        34: "AI_judge is an image-classification project that judges whether an image is real or AI-generated. The README indicates that preprocessing, ResNet18 transfer learning, metric logging, best-model saving, and a Streamlit app are already underway.",
        35: "The current debug focus should be reproducibility: align training code, data paths, model artifact locations, and the app's model-loading path before deeper accuracy work.",
        36: "Debug Summary",
        39: "Design Notes",
        40: "This review treats the current state as a debug inventory rather than a full code rewrite. The five main surfaces are data input, training execution, evaluation logs, model saving, and web inference.",
        41: "Open items are kept as follow-up checks, not final conclusions. The real archive layout and the exact training path that produced the current model artifact still need confirmation.",
        43: "Hypothesis",
        44: "If path, import, and artifact-location inconsistencies are resolved first, training results and web inference can be compared under the same assumptions. That should make later accuracy and UI decisions easier.",
        46: "Success Criteria",
        47: "Recommended priority order: 1. blocking import/path issues, 2. model mismatch between training and inference, 3. reproducible evaluation logs, 4. accuracy improvements, 5. UI improvements.",
        49: "Pre-Fix Validation",
        50: "Confirm that README progress matches the current files.",
        51: "Confirm which training code `run.py` calls and which dataset module actually exists.",
        52: "Confirm whether `loss.csv` and the saved model file belong to the same experiment.",
        53: "Confirm the model filename, class order, and checkpoint format expected by the Streamlit app.",
        54: "After fixes, validate in this order: dataloader only, one-epoch training, then app startup.",
        56: "Current Progress and Evidence",
        57: f"The training log shows the best validation accuracy at epoch {loss['best_epoch']} ({loss['best_valid_accuracy']}); epoch 20 ends at {loss['final_valid_accuracy']}. Together with the README, this suggests the baseline model and app shell have reached a working draft.",
        59: "Primary Outcome",
        61: "Primary progress: ResNet-family training logs exist, best-model saving is implemented, and a Streamlit inference screen is present. Primary risk: file names, save locations, and import names are not yet fully aligned.",
        63: "Guardrail Checks",
        65: "Guardrails for the next fix pass: do not overwrite existing training outputs, do not revert the user's current `model/My_CNN/dataset.py` changes, and do not mix model-comparison conditions.",
        66: "Debug Area Results",
        67: "The review separates dataset handling, baseline training, My_CNN training, evaluation graphing, and the Streamlit app. Keeping these areas separate makes the next repair order clearer.",
        69: "The highest-value first fixes are import and path issues because they can block execution immediately. Accuracy and UI work are easier once the same artifacts can be read and written consistently.",
        71: "Data Quality and Limitations",
        72: "`archive` exists, but this draft does not yet verify the full dataset contents.",
        73: "`loss.csv` is present, but its exact model definition and data split provenance still need confirmation.",
        74: "`model/baseline/ResNet50.py` imports `model.baseline.tensor`, while the visible source file is `dataset.py`.",
        75: "`model/My_CNN/NN_part1.py` also imports `model.My_CNN.tensor`; that source-name mapping needs confirmation.",
        76: "The Streamlit app reads `resnet18_real_fake.pth` from the repository root, which differs from the ResNet50 output path.",
        77: "Therefore, this is a configuration-based debug draft, not a final diagnosis with fresh execution logs.",
        79: "Interpretation",
        80: "The project does not look blocked so much as split across several experiment paths. Aligning references and artifacts should make the next improvement pass much easier.",
        81: "Recommended sequence: fix import/path issues, align trained-model artifacts with the app loader, then proceed to ResNet50/EfficientNet comparisons and data augmentation.",
        83: "Decision Log",
        86: "Next Debug Actions",
        87: "Choose whether dataset imports should use `tensor` or `dataset`, then make the code consistent.",
        88: "Add a simple existence check for `archive/rvf10k/train` and `archive/rvf10k/valid`.",
        89: "Choose one model artifact location: repository root or `output/`.",
        90: "Fix `graph.py` so `loss.csv` is resolved independently of the current working directory.",
        91: "Use one-epoch training and Streamlit startup as the post-fix acceptance checks.",
        93: "Appendix A: Debug Metric Definitions",
        94: "Executability: imports, data paths, and model-save locations are aligned and can be rerun with the same command.",
        95: "Reproducibility: `loss.csv`, model files, and `class_to_idx` are saved as one experiment unit.",
        96: "Inference consistency: the app uses the same class order, preprocessing, and model structure as training.",
        97: "Accuracy-improvement readiness: augmentation, fine-tuning, and model comparisons can be tested under the same evaluation setup.",
        98: "UI readiness: missing-model guidance, upload flow, confidence display, and caveats are clear to users.",
        100: "Appendix B: Analyst Notes",
        101: f"Reference template SHA-256: {reference_hash}",
        102: "This version is a debug-structure sample before detailed logs are added.",
        103: "A later version can add commands, full error text, patches, and confirmation results.",
        104: "Current user changes should be preserved; fixes should be scoped and separated.",
    }

    for idx, text in replacements.items():
        set_paragraph_text(doc.paragraphs[idx], text)

    fill_table(
        doc.tables[0],
        [
            ["Version", "0.1 draft"],
            ["Prepared By", "Riku / Codex"],
            ["Reviewers", "TBD"],
            ["Date Prepared", today],
            ["Reporting Window", "2026-07-14 時点の作業状態"],
            ["Status", "Draft / pre-detail"],
        ],
    )

    fill_table(
        doc.tables[1],
        [
            ["Field", "Details"],
            ["Project Name", "AI_judge"],
            ["Debug Key", "Path / import / model artifact consistency"],
            ["Owner Team", "Personal project"],
            ["Business Owner", "Riku"],
            ["Product Surface", "Python training pipeline + Streamlit app"],
            ["Primary Objective", "Make the next repair order clear"],
            ["Current Implementation", "README, logs, Streamlit app, and ResNet code exist"],
            ["Target State", "Rerunnable with the same data, save path, and model definition"],
            ["Allocation", "該当なし"],
            ["Unit of Randomization", "N/A"],
            ["Audience", "Developer / later reviewer"],
            ["Exclusions", "Detailed logs, external eval data, and unrun tests"],
            ["Start Date", "TBD"],
            ["End Date", "Open"],
            ["Planned Runtime", "smoke check -> 1 epoch -> normal training"],
            ["Actual Runtime", "TBD"],
        ],
    )

    fill_table(
        doc.tables[2],
        [
            ["Metric", "Target/Rule"],
            ["Primary Metric", "`run.py`, dataloaders, and app run under one consistent setup"],
            ["Guardrail 1", "Do not revert current `model/My_CNN/dataset.py` work"],
            ["Guardrail 2", "Do not overwrite existing model and log artifacts"],
            ["Guardrail 3", "Record model, split, and preprocessing for comparisons"],
            ["Decision Rule", "Fix blockers first; improve accuracy afterward"],
        ],
    )

    fill_table(
        doc.tables[3],
        [
            ["Metric", "Current", "Target / Next Check"],
            ["Training log", f"best valid acc {loss['best_valid_accuracy']} at epoch {loss['best_epoch']}", "生成元とモデル保存物を紐づける"],
            ["Final epoch", f"train acc {loss['final_train_accuracy']} / valid acc {loss['final_valid_accuracy']}", "再実行して同じ傾向か見る"],
            ["Model artifact", "`resnet18_real_fake.pth` in root", "Align save path and app load path"],
            ["Data path", "`archive` exists", "Check `train/valid` before execution"],
            ["Smoke check", "Not run", "Check one epoch and Streamlit startup"],
        ],
    )

    fill_table(
        doc.tables[4],
        [
            ["Measure", "Current", "Target", "Gap", "Priority"],
            ["Primary outcome", "Inventory complete", "Rerunnable execution path", "import/path mismatch", "High"],
            ["95% Confidence Interval", "N/A", "N/A", "N/A", "N/A"],
            ["Two-Sided p-value", "N/A", "N/A", "N/A", "N/A"],
            ["Outcome Narrative", "Training and app pieces exist", "Read the same artifact", "Save path/name", "High"],
        ],
    )

    fill_table(
        doc.tables[5],
        [
            ["Guardrail Metric", "Current", "Target", "Delta", "Status"],
            ["User changes", "`dataset.py` modified", "Preserve current work", "Care needed", "Watch"],
            ["Model overwrite risk", "Existing pth file", "Save explicitly or under new name", "Policy needed", "Watch"],
            ["Data availability", "`archive` exists", "Structure checked", "Unverified", "Open"],
            ["App consistency", "Loads ResNet18 pth", "Matches training code", "Needs check", "Open"],
        ],
    )

    fill_table(
        doc.tables[6],
        [
            ["Area", "Current Value", "Target Value", "Gap", "Interpretation"],
            ["README", "Progress list exists", "Synced with implementation", "May be partly stale", "Update candidate"],
            ["baseline", "ResNet50 code exists", "Valid dataset import", "`tensor` name needs check", "Fix early"],
            ["My_CNN", "Custom CNN code exists", "Valid dataset import", "`tensor` name needs check", "Fix early"],
            ["graph.py", "Reads loss.csv", "Stable root-based path", "BASE_DIR may be too deep", "Fix candidate"],
            ["app.py", "Streamlit screen exists", "Model/class order match training", "Artifact origin check", "Check candidate"],
        ],
    )

    fill_table(
        doc.tables[7],
        [
            ["Item", "Decision"],
            ["Final Decision", "Prioritize execution-path stability before accuracy work"],
            ["Decision Date", today],
            ["Approvers", "Riku"],
            ["Rollout Type", "Small fixes; verify dataloader -> training -> app"],
            ["Rollback Trigger", "Stop if a change could overwrite existing models or logs"],
            ["Follow Up", "Add detailed logs, then split into concrete fixes"],
        ],
    )

    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            if "Report Name" in paragraph.text:
                set_paragraph_text(paragraph, paragraph.text.replace("Report Name", "AI_judge Debug Progress"))

    doc.save(OUT)
    replace_package_text(OUT, "Report Name", "AI_judge Debug Progress")
    if sha256(REFERENCE) != reference_hash:
        raise RuntimeError("Reference template changed during build")
    print(OUT)


if __name__ == "__main__":
    main()
